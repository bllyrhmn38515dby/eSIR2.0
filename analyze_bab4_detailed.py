import docx

# Read the document
doc = docx.Document('(SKRIPSI) BillyRahmansyah_15210008_PENERAPAN METODE USER-CENTERED DESIGN (UCD) MENGGUNAKAN WIREFRAMING UNTUK PENINGKATAN NAVIGASI DAN PELACAKAN REAL-TIME PADA APLIKASI SISTEM RUJUKAN ONLINE RU.docx')

print('=== ANALISIS STRUKTUR BAB IV ===')
print('Total paragraphs:', len(doc.paragraphs))

bab4_found = False
sub_sections = []
current_section = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    # Check if we found BAB IV
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
        print(f'\nFound BAB IV at paragraph {i+1}: {text}')
    
    # If we're in BAB IV, look for sub-sections
    elif bab4_found:
        if 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            print(f'\nEnd of BAB IV at paragraph {i+1}: {text}')
            break
        
        # Look for numbered sections (4.1, 4.2, etc.)
        if text.startswith('4.1') or text.startswith('4.2') or text.startswith('4.3') or text.startswith('4.4') or text.startswith('4.5'):
            print(f'Sub-section: {text}')
            sub_sections.append(text)
            current_section = text
        
        # Look for subsections (4.1.1, 4.1.2, etc.)
        elif text.startswith('4.1.') or text.startswith('4.2.') or text.startswith('4.3.') or text.startswith('4.4.') or text.startswith('4.5.'):
            print(f'  Sub-subsection: {text}')
            sub_sections.append(f'  {text}')

print('\n=== STRUKTUR BAB IV YANG DITEMUKAN ===')
for section in sub_sections:
    print(section)

# Also look for any tables or figures mentioned
print('\n=== MENCARI TABEL DAN GAMBAR YANG ADA ===')
bab4_found = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
    
    elif bab4_found:
        if 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            break
        
        # Look for table or figure references
        if 'Tabel' in text or 'tabel' in text or 'Gambar' in text or 'gambar' in text:
            print(f'Found: {text}')

