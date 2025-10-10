import docx

# Read the updated document
doc = docx.Document('(SKRIPSI) BillyRahmansyah_15210008_PENERAPAN METODE USER-CENTERED DESIGN (UCD) MENGGUNAKAN WIREFRAMING UNTUK PENINGKATAN NAVIGASI DAN PELACAKAN REAL-TIME PADA APLIKASI SISTEM RUJUKAN ONLINE RU.docx')

print('=== ANALISIS DETAIL STRUKTUR BAB IV ===')

bab4_found = False
sections = []
current_main_section = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
        print(f'\nFound BAB IV at paragraph {i+1}: {text}')
    
    elif bab4_found:
        if 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            print(f'\nEnd of BAB IV at paragraph {i+1}: {text}')
            break
        
        # Look for main sections
        if text.startswith('4.1') or text.startswith('4.2') or text.startswith('4.3') or text.startswith('4.4') or text.startswith('4.5') or text.startswith('4.6') or text.startswith('4.7'):
            print(f'\nMain Section: {text}')
            sections.append(text)
            current_main_section = text
        
        # Look for subsections
        elif text.startswith('4.1.') or text.startswith('4.2.') or text.startswith('4.3.') or text.startswith('4.4.') or text.startswith('4.5.') or text.startswith('4.6.') or text.startswith('4.7.'):
            print(f'  Subsection: {text}')
            sections.append(f'  {text}')

print('\n=== STRUKTUR LENGKAP BAB IV ===')
for section in sections:
    print(section)

# Look for specific UCD methodology content
print('\n=== MENCARI TAHAP-TAHAP UCD ===')
bab4_found = False
ucd_steps = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
    
    elif bab4_found:
        if 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            break
        
        # Look for UCD methodology steps
        if any(step in text.lower() for step in ['analisis kebutuhan', 'identifikasi pengguna', 'persona', 'user journey', 'wireframe', 'prototype', 'usability testing', 'evaluasi']):
            if len(text) > 30:  # Meaningful content
                ucd_steps.append(f'UCD Step: {text[:200]}...')

print('\nUCD Methodology Steps Found:')
for step in ucd_steps[:8]:  # Show first 8 steps
    print(step)
