import docx

# Read the document
doc = docx.Document('(SKRIPSI) BillyRahmansyah_15210008_PENERAPAN METODE USER-CENTERED DESIGN (UCD) MENGGUNAKAN WIREFRAMING UNTUK PENINGKATAN NAVIGASI DAN PELACAKAN REAL-TIME PADA APLIKASI SISTEM RUJUKAN ONLINE RU.docx')

print('Total paragraphs:', len(doc.paragraphs))
print('\nSearching for BAB IV structure...')

bab4_found = False
bab4_content = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    # Check if we found BAB IV
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
        print(f'Found BAB IV at paragraph {i+1}: {text}')
        bab4_content.append(f'BAB IV: {text}')
    
    # If we're in BAB IV, collect sub-sections
    elif bab4_found:
        if text.startswith('4.') or text.startswith('4 '):
            print(f'Sub-bab: {text}')
            bab4_content.append(f'  {text}')
        elif 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            print(f'End of BAB IV at paragraph {i+1}: {text}')
            break
        elif text and len(text) > 10:  # Collect some content paragraphs
            bab4_content.append(f'    Content: {text[:100]}...')

print('\n=== BAB IV STRUCTURE FOUND ===')
for content in bab4_content[:20]:  # Show first 20 items
    print(content)

