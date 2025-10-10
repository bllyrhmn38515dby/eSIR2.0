import docx

# Read the updated document
doc = docx.Document('(SKRIPSI) BillyRahmansyah_15210008_PENERAPAN METODE USER-CENTERED DESIGN (UCD) MENGGUNAKAN WIREFRAMING UNTUK PENINGKATAN NAVIGASI DAN PELACAKAN REAL-TIME PADA APLIKASI SISTEM RUJUKAN ONLINE RU.docx')

print('=== ANALISIS STRUKTUR BAB IV YANG DIPERBARUI ===')
print('Total paragraphs:', len(doc.paragraphs))

bab4_found = False
bab4_content = []
current_section = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    # Check if we found BAB IV
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
        print(f'\nFound BAB IV at paragraph {i+1}: {text}')
        bab4_content.append(f'BAB IV: {text}')
    
    # If we're in BAB IV, collect structure
    elif bab4_found:
        if 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            print(f'\nEnd of BAB IV at paragraph {i+1}: {text}')
            break
        
        # Look for main sections (4.1, 4.2, etc.)
        if text.startswith('4.1') or text.startswith('4.2') or text.startswith('4.3') or text.startswith('4.4') or text.startswith('4.5') or text.startswith('4.6') or text.startswith('4.7'):
            print(f'Main section: {text}')
            bab4_content.append(f'  {text}')
            current_section = text
        
        # Look for subsections (4.1.1, 4.1.2, etc.)
        elif text.startswith('4.1.') or text.startswith('4.2.') or text.startswith('4.3.') or text.startswith('4.4.') or text.startswith('4.5.') or text.startswith('4.6.') or text.startswith('4.7.'):
            print(f'    Subsection: {text}')
            bab4_content.append(f'    {text}')
        
        # Look for UCD-related content
        elif 'UCD' in text.upper() or 'User-Centered Design' in text or 'user-centered' in text.lower():
            print(f'UCD Content: {text[:100]}...')
            bab4_content.append(f'      UCD: {text[:100]}...')

print('\n=== STRUKTUR BAB IV YANG DITEMUKAN ===')
for content in bab4_content[:30]:  # Show first 30 items
    print(content)

# Look specifically for UCD process content
print('\n=== MENCARI KONTEN PROSES UCD ===')
bab4_found = False
ucd_content = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    
    if 'BAB IV' in text.upper() or 'BAB 4' in text.upper():
        bab4_found = True
    
    elif bab4_found:
        if 'BAB V' in text.upper() or 'BAB 5' in text.upper():
            break
        
        # Look for UCD process steps
        if any(keyword in text.lower() for keyword in ['ucd', 'user-centered', 'persona', 'wireframe', 'prototype', 'usability', 'testing']):
            if len(text) > 20:  # Only meaningful content
                ucd_content.append(f'UCD Content: {text[:150]}...')

print('\nUCD-related content found:')
for content in ucd_content[:10]:  # Show first 10 UCD content
    print(content)
